import docx,os,pyperclip,sys,time
from tkinter import*
import tkinter.messagebox


root = Tk()
root.iconbitmap(r'C:\Users\HARI\AppData\Local\Programs\Python\Python35-32\DLLs\My_Icon.ico')
root.wm_title("Hari Word Manager")
tkinter.messagebox.showinfo("Hari Technologies","Welcome to HARI Word MANAGER.Converter\n->User-Friendly\n\nFOLLOWING POINTS MAY BE NOTED:\n1.Hyperlinks are just Addresses to a Website,They are no longer a text...So they are not Conerted into .txt\n2.Cliparts,Tables,Charts and all other Visual Information will not be converted(Because they are not Supported by .txt)\n3.To increase Flexibility of .txt files we are not Converting BOLD and ITALIC styles,and also Fonts as they are not Supported like Leading Text Editors like\nNOTEPAD++\n\n\nThanks For Chosing us!!")
root.destroy()
root.mainloop()
print("Enter Addres of the Word File to Handle with.")
time.sleep(1)

ClipBoard = pyperclip.paste()
if(ClipBoard != "" and ((".docx" in ClipBoard )or( ".doc" in ClipBoard))):
    obj = Tk()
    obj.iconbitmap(r'C:\Users\HARI\AppDAta\Local\Programs\Python\Python35-32\DLLs\My_Icon.ico')
    obj.wm_title("Hari Word to TXT Converter")
    Answer = tkinter.messagebox.askquestion("Hari Technologies","We have Detected Text in your Clip Board\n"+ClipBoard+"\nDo you Want to use it?")
    obj.destroy()
    obj.mainloop()
    if(Answer == "yes"):
        WordAddress = ClipBoard
    if(Answer == "no"):
        WordAddress = input()
else:
    WordAddress=input()
root.mainloop()


try:
    Doc_obj = docx.Document(WordAddress)
except PackageNotFoundError:
    root = Tk()
    root.iconbitmap(r'C:\Users\HARI\AppData\Local\Programs\Python\Python35-32\DLLs\My_Icon.ico')
    root.wm_title("Hari Excell Manager")
    tkinter.messagebox.showinfo("Hari Technologies"," Package Not Found...\n\nClick OK to Abort Program")
    root.destroy()
    root.mainloop()
    sys.exit(0)

TxtFileName = os.path.relpath(WordAddress , os.path.dirname(WordAddress))
TxtFileName = TxtFileName[:-5]
TxtFileName = TxtFileName+" _"+os.getlogin()+".txt"
FileAddress = os.path.dirname(WordAddress)+"//"+TxtFileName
File = open(FileAddress , "w")

Doc_obj = docx.Document(WordAddress)

for i in range(0 , len(Doc_obj.paragraphs)):
    try:
        File.write(Doc_obj.paragraphs[i].text)
        File.write("\n\n")
    except UnicodeEncodeError:
        root = Tk()
        root.iconbitmap(r'C:\Users\HARI\AppData\Local\Programs\Python\Python35-32\DLLs\My_Icon.ico')
        root.wm_title("Hari Word Manager")
        tkinter.messagebox.showinfo("Hari Technologies","This Word File Contains Content that tis not supported by .txt\nEg:Special Symbols\nERROR(Technical):Encoding UTF-8 PRoblem\n\nClick OK to Abort!!")
        root.destroy()
        root.mainloop()
        File.close()
        File = open(FileAddress , "w")
        File.write("UnicodeError Detected..Please Verify the Parent Word Doccument")
        File.close()
        sys.exit()

File.close()




root = Tk()
root.iconbitmap(r'C:\Users\HARI\AppData\Local\Programs\Python\Python35-32\DLLs\My_Icon.ico')
root.wm_title("Hari Word Manager")
tkinter.messagebox.showinfo("Hari Technologies","File Converted Successfully!!!!!!\n\nThanks For Chosing us!!")
root.destroy()
root.mainloop()


