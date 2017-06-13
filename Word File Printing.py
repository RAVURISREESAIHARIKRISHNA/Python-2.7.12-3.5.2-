import docx

print("Enter the Location of the Word Document to Display the Text")
FileAddress = input()

Doc_obj = docx.Document(FileAddress)
Total_Paras = len(Doc_obj.paragraphs)
for i in range(0,Total_Paras):
	print(Doc_obj.paragraphs[i].text)
	print("\n")
